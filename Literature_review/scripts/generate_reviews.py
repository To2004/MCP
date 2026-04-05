"""Generate comprehensive Datasets and Benchmarks review files from MCP security literature.

Outputs 6 files: datasets_review.md/.docx/.xlsx and benchmarks_review.md/.docx/.xlsx
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

OUTPUT_DIR = Path(__file__).parent
TODAY = datetime.date.today().isoformat()

# ---------------------------------------------------------------------------
# SECTION 1: DATASETS DATA
# ---------------------------------------------------------------------------

DATASETS: list[dict[str, str]] = [
    # ── MCP-Specific Datasets ──
    {
        "dataset_name": "MCP-AttackBench",
        "paper_title": "MCP-Guard: A Multi-Stage Defense-in-Depth Framework for Securing Model Context Protocol in Agentic AI (Xing et al., 2025)",
        "link": "https://arxiv.org/abs/2508.10991",
        "description": "A large-scale synthetic security dataset covering unique threat vectors for MCP environments. It uses a hierarchical taxonomy organized into three families: Semantic & Adversarial attacks, Protocol-Specific attacks, and Injection & Execution attacks. Designed to train and evaluate cascaded defense pipelines for MCP security.",
        "structure": "70,448 samples total. Training corpus of 5,258 samples (2,153 adversarial, 3,105 benign) used for fine-tuning E5 embedding model. Covers three attack families with hierarchical sub-categories. Format: structured text samples with attack/benign labels.",
        "paper_usage": "Used as the primary training and evaluation dataset for the MCP-Guard three-stage cascaded defense pipeline. The training corpus fine-tuned the Multilingual-E5-large embedding model (Stage II neural detection), achieving 96.01% F1 score with a 95.06% improvement over baseline.",
        "project_value": "Directly applicable for training the risk_scorer module of an MCP security system. The hierarchical attack taxonomy can inform severity tier definitions. The 70K+ labeled samples provide substantial training data for ML-based risk classification of MCP tool invocations.",
    },
    {
        "dataset_name": "MCPTox Dataset",
        "paper_title": "MCPTox: A Benchmark for Tool Poisoning Attack on Real-World MCP Servers (Wang et al., 2025)",
        "link": "https://openreview.net/forum?id=xbs5dVGUQ8",
        "description": "The first benchmark built from real-world MCP servers for evaluating tool poisoning attacks. Contains malicious test cases across three attack paradigms: Explicit Trigger Function Hijacking, Implicit Trigger Function Hijacking, and Implicit Trigger Parameter Tampering. Covers 10 attack categories including shadowing, puppet, and spoofing attacks.",
        "structure": "45 live real-world MCP servers, 353 authentic tools, 1,312-1,497 malicious test cases. Breakdown: Explicit Trigger—Function Hijacking (224 cases), Implicit Trigger—Function Hijacking (548 cases), Implicit Trigger—Parameter Tampering (725 cases). Format: structured test cases with user queries + poisoned tool descriptions.",
        "paper_usage": "Primary evaluation benchmark for measuring tool poisoning attack success rates across multiple LLM agents (GPT-4o, Qwen3-32b, Claude, etc.). Found overall 72.8% attack success rate with significant model vulnerability variation.",
        "project_value": "Essential for evaluating an MCP risk scoring system's ability to detect tool poisoning attacks. The three attack paradigms map to different severity tiers. Real-world server data ensures the risk scorer is tested against realistic threats rather than synthetic patterns only.",
    },
    {
        "dataset_name": "MCPSecBench Playground",
        "paper_title": "MCPSecBench: A Systematic Security Benchmark and Playground for Testing Model Context Protocols (Yang et al., 2025)",
        "link": "https://arxiv.org/abs/2508.13220",
        "description": "A comprehensive MCP security testing framework providing taxonomy-driven test cases, a GUI test harness, prompt datasets, attack scripts, and server/client implementations. Covers 17 distinct attack types across 4 attack surfaces: client-side, protocol-side, server-side, and host-side.",
        "structure": "17 attack types across 4 attack surfaces. Prompt dataset with LLM-based prompt variation. Multi-layer evaluation covering Claude Desktop, OpenAI GPT-4.1, and Cursor v2.3.29. 15 trials per attack for statistical robustness. Cost: $0.41-$0.76 per testing round.",
        "paper_usage": "Used as both a benchmark and playground for evaluating MCP security across multiple platforms and providers. Found 100% ASR on protocol-side attacks (universal), ~33% on client-side, ~47% on server-side, and ~27% on host-side.",
        "project_value": "The 4-surface attack taxonomy directly maps to the risk scoring system's evaluation dimensions. The playground environment enables reproducible testing of risk classification accuracy. Protocol-side attacks (100% ASR) highlight critical areas where the risk scorer must flag maximum severity.",
    },
    {
        "dataset_name": "MCP-SafetyBench Dataset",
        "paper_title": "MCP-SafetyBench: A Benchmark for Safety Evaluation of Large Language Models with Real-World MCP Servers (Zong et al., 2025)",
        "link": "https://arxiv.org/abs/2512.15163",
        "description": "A benchmark evaluating LLM safety across five real-world MCP domains with 20 attack types using multi-turn interaction evaluation. Covers both stealth and disruption attacks across server-side, host-side, and user-side categories using real-world MCP servers with adversarially modified tool descriptions.",
        "structure": "5 domains (Location Navigation, Repository Management, Financial Analysis, Browser Automation, Web Search). 20 attack types. Evaluated on 13 LLMs (9 proprietary including GPT-5, Claude-4.0-Sonnet, Grok-4; 4 open-source). Multi-turn interaction format.",
        "paper_usage": "Primary benchmark evaluating 13 LLMs on safety compliance across all domain-attack combinations. Measured Task Success Rate (TSR), Attack Success Rate (ASR), and safety prompt effectiveness.",
        "project_value": "The 5-domain coverage ensures the risk scoring system generalizes across MCP use cases. The 20 attack types provide a comprehensive threat catalog for severity classification. Multi-turn evaluation tests whether risk scoring remains accurate across extended agent sessions.",
    },
    {
        "dataset_name": "MCIP-Bench",
        "paper_title": "MCIP: Protecting MCP Safety via Model Contextual Integrity Protocol (Jing et al., 2025)",
        "link": "https://github.com/HKUST-KnowComp/MCIP",
        "description": "A benchmark suite for evaluating LLMs' safety capabilities in MCP interactions. Covers 11 categories: 10 risk types (Identity Injection, Function Overlapping, Function Injection, Data Injection, Excessive Privileges Overlapping, Function Dependency Injection, Replay Injection, Wrong Parameter Intent Injection, Ignore Purpose Intent Injection, Causal Dependency Injection) plus 1 safe/gold class.",
        "structure": "2,218 synthesized instances total (1,192 from Glaive AI + 1,026 from ToolACE). Each instance averages ~6 dialogue turns. 11 categories. Format: multi-turn function-calling dialogues with safety labels.",
        "paper_usage": "Used to evaluate MCIP Guardian and baseline models (xLAM, ToolACE, Qwen2.5, DeepSeek-R1) on Safety Awareness (binary classification) and Risk Resistance (11-class identification). Measured by Accuracy and Macro-F1.",
        "project_value": "The 10 risk types provide a ready-made taxonomy for the risk scoring system's severity classification. The structured dialogue format tests multi-turn risk assessment. Can directly evaluate whether the MCP risk scorer correctly identifies each of the 10 risk categories.",
    },
    {
        "dataset_name": "MCIP Guardian Training Dataset",
        "paper_title": "MCIP: Protecting MCP Safety via Model Contextual Integrity Protocol (Jing et al., 2025)",
        "link": "https://github.com/HKUST-KnowComp/MCIP",
        "description": "A training dataset in Model Contextual Integrity (MCI) format for training safety-aware MCP models. Covers all 11 MCIP categories with structured information transmission steps.",
        "structure": "13,830 instances across 11 categories. Distribution: True (1,791), Identity Injection (1,749), Function Overlapping (1,395), Function Injection (1,382), Data Injection (1,361), Excessive Privileges (1,401), Function Dependency Injection (1,372), Replay Injection (1,371), Wrong Parameter (664), Ignore Purpose (718), Causal Dependency (626). ~8 transmission steps per instance.",
        "paper_usage": "Used to fine-tune Salesforce/Llama-xLAM-2-8b-fc-r as the MCIP Guardian safety model. The fine-tuned model significantly outperformed baseline LLMs on safety awareness tasks.",
        "project_value": "Directly usable for training the risk_scorer module to classify MCP interactions into 10 risk categories. The imbalanced class distribution reflects real-world threat frequency patterns. The MCI format (information transmission steps) aligns with how MCP tool invocations flow through the protocol.",
    },
    {
        "dataset_name": "ProtoAMP Attack Data",
        "paper_title": "Breaking the Protocol: Security Analysis of the MCP Specification and Prompt Injection Vulnerabilities (Maloyan & Namiot, 2026)",
        "link": "https://arxiv.org/abs/2601.17549",
        "description": "Protocol Amplification benchmark dataset containing controlled test scenarios measuring how MCP architecture amplifies prompt injection attack success rates compared to non-MCP baselines. Attacks injected at three protocol layers: resource content, tool response payloads, and sampling request prompts.",
        "structure": "847 controlled test scenarios. Tested on 3 MCP server implementations (filesystem, git, sqlite) and 4 LLM backends (Claude-3.5-Sonnet, GPT-4o, Llama-3.1-70B). Three injection layers measured independently.",
        "paper_usage": "Quantified that MCP architecture amplifies attack success by +23-41% versus non-MCP baselines. Indirect injection: 47.8% ASR vs 31.2% baseline (+16.6%). Cross-server propagation: 61.3% vs 19.7% (+41.6%). Sampling vulnerability: 67.2% ASR.",
        "project_value": "Demonstrates that MCP-specific protocol features create amplified risk — critical evidence for the risk scorer to assign higher severity to protocol-level attack vectors. The three injection layers map to distinct risk dimensions the scoring system must evaluate.",
    },
    {
        "dataset_name": "MCP Server Dataset (67,057 servers)",
        "paper_title": "Toward Understanding Security Issues in the Model Context Protocol Ecosystem (Li & Gao, 2025)",
        "link": "https://arxiv.org/abs/2510.16558",
        "description": "Large-scale dataset of MCP servers collected across 6 registries for ecosystem-wide security analysis. Covers server metadata, tool definitions, and security posture indicators including invalid links, empty content, missing documentation, and credential leakage.",
        "structure": "67,057 servers across 6 registries: mcp.so (15,000+), MCP Market (10,000+), MCP Store (5,000+), Pulse MCP (7,000+), Smithery (7,682), npm (52,102). 44,499 Python-based tools extracted via AST analysis. 52.1% Python, 22.5% JavaScript, 11.6% TypeScript.",
        "paper_usage": "Used for ecosystem-wide trust analysis identifying: tool confusion attacks (20-100% success), tool shadowing (40-100% success), credential leakage (9 PATs, 3 API keys), server hijacking (111+ instances), and affix-squatting (408 groups).",
        "project_value": "Provides the scale data needed to understand MCP ecosystem risk distribution. The vulnerability prevalence statistics (e.g., 6.75% invalid links, credential leakage rates) can calibrate the risk scorer's base rates. The 44K+ tool database enables training a tool-level risk classifier.",
    },
    {
        "dataset_name": "MCP Server Database (1,360 servers / 12,230 tools)",
        "paper_title": "Mind Your Server: A Systematic Study of Parasitic Toolchain Attacks on the MCP Ecosystem (Zhao et al., 2025)",
        "link": "https://arxiv.org/abs/2509.06572",
        "description": "Database of MCP servers and tools analyzed for parasitic toolchain attack (MCP-UPD) vulnerability. Categorizes tools into External Ingestion Tools (EIT), Privacy Access Tools (PAT), and Network Access Tools (NAT) based on their exploit-enabling capabilities.",
        "structure": "1,360 servers from 3 sources: Pulse MCP (784), MCP Market (310), Awesome MCP Servers (266). 12,230 tools total. Risk categories: EIT (472 tools, 128 servers), PAT (391 tools, 155 servers), NAT (180 tools, 89 servers). 1,062 tools (8.7%) identified with exploit risk.",
        "paper_usage": "Used for ecosystem census revealing 27.2% of servers expose exploitable tool combinations. Demonstrated 90% success rate in constructing real-world parasitic toolchain attacks across 10 toolchains.",
        "project_value": "The EIT/PAT/NAT risk taxonomy provides a tool-level classification scheme for the risk scorer. The finding that 8.7% of tools carry exploit risk gives a baseline detection target. Parasitic toolchain patterns inform multi-tool risk aggregation logic.",
    },
    {
        "dataset_name": "MCP Server Registry Dataset (1,899 repos)",
        "paper_title": "Model Context Protocol (MCP): Landscape, Security Threats, and Future Research Directions (Hou et al., 2025)",
        "link": "https://arxiv.org/abs/2503.23278",
        "description": "Collection of open-source MCP server repositories from 6 registries used to map the MCP ecosystem landscape, server lifecycle, and security posture. Provides the first comprehensive census of the MCP server ecosystem.",
        "structure": "1,899 repositories from: official GitHub repo, Glama, PulseMCP, Smithery, MCP.so, OpenSumi. Includes server metadata, tool schemas, transport configurations, and lifecycle states.",
        "paper_usage": "Used for the first comprehensive MCP landscape analysis, mapping protocol architecture, threat surfaces, and research gaps. Identified key security threats across protocol layers.",
        "project_value": "Provides ecosystem context for the risk scoring system — understanding server population distribution, common tool patterns, and typical configurations helps calibrate risk baselines and identify outlier behaviors worth flagging.",
    },
    {
        "dataset_name": "MCP Server Empirical Dataset (1,899 repos)",
        "paper_title": "Model Context Protocol (MCP) at First Glance: Studying the Security and Maintainability of MCP Servers (Hasan et al., 2025)",
        "link": "https://github.com/SAILResearch/replication-25-mcp-server-empirical-study",
        "description": "Large-scale empirical study dataset of open-source MCP servers examining security vulnerabilities, code quality, and maintainability. Analyzed using SonarQube for vulnerability detection and CHAOSS metrics for ecosystem health.",
        "structure": "1,899 repositories from official, community, and mined (GitHub SDK import search) sources. Analyzed with SonarQube (vulnerability/code smell/bug detection), mcp-scan (tool poisoning), and 14 CHAOSS OSS health metrics.",
        "paper_usage": "First large-scale empirical study identifying common vulnerability patterns across MCP server implementations, dependency auditing results, and ecosystem quality metrics for nearly 1,900 servers.",
        "project_value": "The vulnerability pattern data from 1,900 servers can train the risk scorer to recognize common security anti-patterns. SonarQube findings provide ground truth for code-level risk indicators. Maintainability metrics could serve as proxy signals for server trustworthiness.",
    },
    {
        "dataset_name": "MCP API Usage Dataset (2,117 repos)",
        "paper_title": "We Urgently Need Privilege Management in MCP: A Measurement of API Usage in MCP Ecosystems (Li et al., 2025)",
        "link": "https://arxiv.org/abs/2507.06250",
        "description": "Measurement dataset of API usage across real MCP applications, documenting privilege-sensitive API exposure and over-privilege patterns. Collected from 3 registries to demonstrate the urgent need for least-privilege enforcement.",
        "structure": "2,117 unique GitHub repositories from mcp.so, Glama, and Smithery. Contains tool/API definitions, over-privilege rate measurements, and sensitive API exposure classification.",
        "paper_usage": "Demonstrated that MCP tools routinely access privilege-sensitive APIs without proper authorization controls. Provided evidence-based argument for least-privilege enforcement in MCP ecosystems.",
        "project_value": "The over-privilege measurements directly inform how the risk scorer should weight API access patterns. Privilege-sensitivity classifications can be used as features in the risk scoring model. Evidence of widespread over-privilege justifies strict default-deny policies.",
    },
    {
        "dataset_name": "MCP Ecosystem Dataset (8,401 projects)",
        "paper_title": "A Measurement Study of Model Context Protocol Ecosystem (Guo et al., 2025)",
        "link": "https://arxiv.org/abs/2509.25292",
        "description": "Characterization dataset of the MCP ecosystem covering servers, tools, and registries across 6 major marketplaces. Analyzes growth patterns, distribution, and observable risk indicators.",
        "structure": "8,401 valid MCP projects across 6 marketplaces. Covers server types, tool categories, transport protocols, language distributions, and value classifications.",
        "paper_usage": "Used for ecosystem-scale characterization analyzing distribution of server types, tool categories, growth patterns, and observable risk indicators across the MCP landscape.",
        "project_value": "The ecosystem distribution data helps set risk scoring priors — understanding what 'normal' looks like enables anomaly-based risk detection. Growth pattern data can inform how quickly the risk scorer's training data becomes stale.",
    },
    {
        "dataset_name": "Top-296 MCP Server Dataset",
        "paper_title": "Securing AI Agent Execution — AgentBound (Bühler et al., 2025)",
        "link": "https://github.com/MCP-Security/MCP-Artifact",
        "description": "Dataset of the 296 most popular MCP servers by GitHub stars (59-63,215 stars) from PulseMCP. Used to evaluate automated security policy generation accuracy for the AgentBound access control framework.",
        "structure": "296 MCP servers (top 300 by stars, 296 successfully downloaded). 48 servers manually evaluated (~8 hours each by 2 authors). 816 total permissions across 48 servers (17 avg per server). AgentManifestGen matched 787 permissions (96.5% accuracy).",
        "paper_usage": "Primary evaluation dataset for AgentBound's policy generation. Demonstrated 80.9% accuracy and 100% recall in automated manifest creation. 96 servers evaluated via GitHub Issues.",
        "project_value": "The permission profiles of 296 popular servers provide reference data for what 'normal' access patterns look like. The 17-permissions-per-server average establishes a baseline for the risk scorer to flag abnormal permission requests.",
    },
    {
        "dataset_name": "Malicious MCP Server Dataset (Song et al.)",
        "paper_title": "Beyond the Protocol: Unveiling Attack Vectors in the MCP Ecosystem (Song et al., 2025)",
        "link": "https://github.com/MCP-Security/MCP-Artifact",
        "description": "Public dataset of 4 MCP servers with known categories of malicious behaviors: Google Maps Server (malicious external resource), mcp_server_time (puppet attack), mcp-weather-server (dynamic API host rewriting), and wechat-mcp (SQL injection).",
        "structure": "4 malicious MCP servers, each targeting a different attack vector. Ready-to-use package format for security testing.",
        "paper_usage": "Used to test AgentBox's security enforcement (RQ2). AgentBound successfully blocked all 4 attack types through its policy-based access control.",
        "project_value": "Provides labeled ground truth for the risk scorer — 4 known-malicious servers with documented attack vectors for testing detection accuracy. Useful as a quick validation set for the risk scoring pipeline.",
    },
    {
        "dataset_name": "Damn Vulnerable MCP Server",
        "paper_title": "Securing AI Agent Execution — AgentBound (Bühler et al., 2025)",
        "link": "https://github.com/harishsg993010/damn-vulnerable-MCP-server",
        "description": "A public security challenge dataset containing 10 intentionally vulnerable MCP servers designed for security testing, similar to DVWA for web applications. Each server contains one or multiple attack vectors.",
        "structure": "10 vulnerable MCP servers (C.1-C.10). Includes: 2 tool poisoning attacks, 1 rug pull attack, 2 malicious external resource attacks, 2 prompt injection attacks, and additional mixed vectors.",
        "paper_usage": "Used to evaluate AgentBox against a diverse set of security challenges. AgentBox prevented 9 out of 10 attacks (all environment-targeting attacks blocked).",
        "project_value": "An excellent test bed for the risk scoring system — 10 servers with known vulnerabilities at varying severity levels. Can validate that the risk scorer assigns appropriate severity tiers to different attack types. The 'capture the flag' format enables iterative risk scorer improvement.",
    },
    {
        "dataset_name": "Component-based Attack PoC Dataset (132 servers)",
        "paper_title": "When MCP Servers Attack: Taxonomy, Feasibility, and Mitigation (Zhao et al., 2025)",
        "link": "Paper under review — contact: Weibo Zhao et al. via https://arxiv.org/abs/2509.24272",
        "description": "A set of 12 Proof-of-Concept malicious MCP servers (one per attack category A1-A12) plus 120 generated malicious servers from modular component seeds. The generator can theoretically produce up to 1,046,529 unique malicious server configurations.",
        "structure": "12 hand-crafted PoC servers + 120 generated servers (10 per category). Generator seeds: 5 malicious launch commands, 7 initialization snippets, 10 malicious/30 benign tools, 10 malicious/10 benign resources, 5 malicious/5 benign prompts. Evaluated across 3 hosts and 5 LLMs, 15 trials each.",
        "paper_usage": "Primary evaluation for taxonomy study. Also tested scanner effectiveness: mcp-scan detected only 4/120 poisoned servers; AI-Infra-Guard performed better but still insufficient.",
        "project_value": "The component-based server generation approach can produce large-scale training data for the risk scorer. The 12-category taxonomy provides a comprehensive attack classification scheme. Scanner evaluation results establish baselines to improve upon.",
    },
    {
        "dataset_name": "MCP Attack Benchmark (Beyond the Protocol)",
        "paper_title": "Beyond the Protocol: Unveiling Attack Vectors in the Model Context Protocol Ecosystem (Song et al., 2025)",
        "link": "https://github.com/MCP-Security/MCP-Artifact",
        "description": "End-to-end evaluation framework testing malicious MCP server attacks across LLMs and MCP clients. Tests three attack tasks (Privacy Steal, Result Manipulation, Cryptocurrency Theft) against three attack vectors (Tool Poisoning, Puppet Attack, Malicious External Resources).",
        "structure": "Two benchmarks: (1) 3 tasks x 3 vectors x 5 LLMs x 20 tests = 900 tests; (2) 3 tasks x 3 vectors x 5 MCP clients x 20 tests = 900 tests. LLMs: Claude 3.7, GPT-4o, DeepSeek-V3, LLaMA-3.1-70B, Gemini 2.5 Pro. Clients: Cherry Studio, Claude Desktop, Cline, Copilot-MCP, Cursor.",
        "paper_usage": "Found average 53% ASR with <5% refusal rate. Exploitation via Malicious External Resources achieved highest ASR (93.33%). Also included a 20-participant user study on a simulated aggregator platform.",
        "project_value": "The cross-LLM and cross-client evaluation reveals which combinations are most vulnerable — informing how the risk scorer should weight model/client context. The 93% ASR for external resource attacks highlights a critical severity tier.",
    },
    {
        "dataset_name": "MCPShield Evaluation Suite (6 test suites)",
        "paper_title": "MCPShield: A Security Cognition Layer for Adaptive Trust Calibration in MCP Agents (Zhou et al., 2026)",
        "link": "https://arxiv.org/abs/2602.14281",
        "description": "Multi-suite security evaluation framework containing 6 distinct malicious MCP server test suites: MCPSafetyBench (tool poisoning + command injection), MCPSecBench (tool substitution + workflow poisoning), DemonAgent (encoded payloads + persistence), MCP-Artifact (result manipulation), Adaptive Monitor (monitor hijacking), and Rug Pull Attack (temporal drift).",
        "structure": "6 test suites, each containing multiple malicious servers with different attack patterns. Evaluated using Tool Safety Rate (TSR), Detection Success Rate (DSR), Attack Success Rate (ASR), time overhead, and token overhead metrics.",
        "paper_usage": "Comprehensive evaluation of MCPShield's three-phase lifecycle defense: pre-invocation probing, execution isolation, and post-invocation reasoning. Achieved 95.30% defense rate versus 10.05% baseline.",
        "project_value": "The 6 diverse test suites cover the full spectrum of MCP attack patterns. The lifecycle-based evaluation (pre/during/post invocation) directly maps to when the risk scorer should assess threats. The 95.30% defense rate sets a target benchmark for the risk scoring system.",
    },
    {
        "dataset_name": "MCP-ITP Implicit Poisoning Data",
        "paper_title": "MCP-ITP: An Automated Framework for Implicit Tool Poisoning in MCP (Li et al., 2026)",
        "link": "https://arxiv.org/abs/2601.07395",
        "description": "Automated framework-generated dataset of implicitly poisoned tool descriptions that achieve high attack success through innocuous-appearing text optimized via black-box techniques. The poisoned descriptions evade existing detection with only 0.3% detection rate.",
        "structure": "Based on MCPTox's 548 implicit poisoning test cases. Evaluated across 12 LLM agents (o1-mini, GPT-4o-mini, GPT-3.5-turbo, DeepSeek-R1/V3, Gemini-2.5-flash, Qwen3 variants). Detection tested against AI-Infra-Guard and Oracle detectors.",
        "paper_usage": "Demonstrated 84.2% attack success rate with only 0.3% detection by existing defenses. Showed that individual words in tool descriptions can steer agent behavior without explicit malicious instructions.",
        "project_value": "Reveals a critical blind spot for risk scoring: implicit poisoning that evades pattern-based detection. The risk scorer must incorporate semantic analysis beyond keyword matching. The 0.3% detection rate establishes the difficulty baseline for implicit attack detection.",
    },
    {
        "dataset_name": "Log-To-Leak Attack Scenarios",
        "paper_title": "Log-To-Leak: Prompt Injection Attacks on Tool-Using LLM Agents via Model Context Protocol (Hu et al., 2026)",
        "link": "https://openreview.net/forum?id=UVgbFuXPaO",
        "description": "Prompt injection attack framework exploiting MCP logging tools for data exfiltration. Demonstrates that benign infrastructure tools (logging, monitoring) can be weaponized to silently capture and transmit sensitive data while preserving task quality.",
        "structure": "Attack scenarios targeting MCP logging and monitoring tools. Evaluation of leakage success rate, task quality preservation, and detection evasion across MCP configurations.",
        "paper_usage": "Demonstrated successful covert data exfiltration through benign MCP infrastructure tools while maintaining normal task completion quality, making the attack difficult to detect through output monitoring.",
        "project_value": "Highlights that the risk scorer must evaluate seemingly benign tools (logging, monitoring) for covert channel potential. Risk scoring should account for tool combinations, not just individual tools. Task quality preservation during attacks means the risk scorer cannot rely solely on output quality as a safety signal.",
    },
    {
        "dataset_name": "MCPSafetyScanner Test Scenarios",
        "paper_title": "MCP Safety Audit: LLMs with the Model Context Protocol Allow Major Security Exploits (Radosevich & Halloran, 2025)",
        "link": "https://github.com/johnhalloran321/mcpSafetyScanner",
        "description": "The first automated MCP security auditing tool with test scenarios for credential theft, malicious code execution, and coercion attacks. Tested against real MCP deployments using standard MCP servers (filesystem, Slack, Everything, Chroma).",
        "structure": "Qualitative attack demonstrations on 4 real MCP servers. Tested with Claude 3.7 and Llama-3.3-70B-Instruct. Attack types: MCE (Malicious Code Execution), RAC (Remote Access Control), CT (Credential Theft), RADE (Retrieval-Agent Deception).",
        "paper_usage": "Demonstrated coercion attacks enabling credential theft and malicious code execution against real MCP deployments through systematic vulnerability scanning with MCPSafetyScanner.",
        "project_value": "MCPSafetyScanner provides a ready-made tool for validating the risk scorer's detection capabilities. The documented attack types (MCE, RAC, CT, RADE) define critical severity categories. The real-world MCP server test environment ensures practical relevance.",
    },
    # ── Agent Security Datasets ──
    {
        "dataset_name": "InjecAgent Dataset",
        "paper_title": "InjecAgent: Benchmarking Indirect Prompt Injections in Tool-Integrated Large Language Model Agents (Zhan et al., 2024)",
        "link": "https://github.com/uiuc-kang-lab/InjecAgent",
        "description": "Benchmark for indirect prompt injections in tool-integrated LLM agents with explicit study of private-data exfiltration. Provides user tools, attacker tools, and multiple agent configurations across direct harm and data stealing categories.",
        "structure": "1,054 test cases total. 17 user tools and 62 attacker tools. Two harm categories: direct harm and data stealing/exfiltration. Multiple agent configurations for vulnerability assessment.",
        "paper_usage": "First indirect injection benchmark for tool-integrated agents. Systematically evaluated attack success rates across tools, agents, harm categories, and exfiltration scenarios.",
        "project_value": "The user-tool vs attacker-tool distinction maps directly to the risk scorer's need to differentiate legitimate and adversarial tool access. The exfiltration category is directly relevant to MCP data leakage risk assessment. Provides labeled test data for binary risk classification.",
    },
    {
        "dataset_name": "AgentDojo Task Suites",
        "paper_title": "AgentDojo: A Dynamic Environment to Evaluate Prompt Injection Attacks and Defenses for LLM Agents (Debenedetti et al., 2024)",
        "link": "https://github.com/ethz-spylab/agentdojo",
        "description": "Dynamic evaluation environment for prompt injection attacks and defenses in tool-using LLM agents. Provides realistic tasks with parameterized attacks, defense baselines, and reproducible benchmarking across 4 agent environments.",
        "structure": "4 task suites: Workspace, Travel, Banking, Slack. 97 user tasks and 629 injection tasks. Supports multiple attack implementations and defense baselines. Dynamic environment with parameterized attack generation.",
        "paper_usage": "Used as the primary evaluation platform by Progent (41.2% -> 2.2% ASR reduction), LlamaFirewall, TraceAegis, ToolSafe, and The Task Shield. The most widely adopted agent security benchmark.",
        "project_value": "The most established evaluation platform for agent security — essential for benchmarking the MCP risk scoring system against existing defenses. The 4 realistic domains ensure broad coverage. Widely adopted, enabling direct comparison with published results.",
    },
    {
        "dataset_name": "R-Judge Records",
        "paper_title": "R-Judge: Benchmarking Safety Risk Awareness for LLM Agents (Yuan et al., 2024)",
        "link": "https://github.com/Lordog/R-Judge",
        "description": "Benchmark for evaluating risk awareness in LLM agent behaviors across diverse environments. Contains curated multi-turn agent interaction records with annotated safety labels and risk descriptions spanning 27 key risk scenarios.",
        "structure": "569 multi-turn agent interaction records. 27 risk scenarios across 5 application categories. 10 risk types with annotated safety labels. Sources: WebArena, ToolEmu, InterCode-Bash, InterCode-SQL, MINT environments.",
        "paper_usage": "Evaluated 11 LLMs showing GPT-4o achieves 74.42% while others near random. Fine-tuning on safety judgment significantly improves performance.",
        "project_value": "The 10 risk types and 27 scenarios provide a template for defining MCP-specific risk categories. The scoring approach (judging agent interactions) directly mirrors the MCP risk scorer's goal. The 74.42% best accuracy shows the current difficulty level for risk assessment.",
    },
    {
        "dataset_name": "ASB (Agent Safety Benchmark)",
        "paper_title": "Agent Security Bench (Zhang et al., 2025) / Used by ToolSafe (Mou et al., 2026) and Progent (Shi et al., 2025)",
        "link": "https://github.com/agiresearch/ASB",
        "description": "Benchmark for formalizing and evaluating attacks and defenses in LLM-based agents. Contains multiple attack prompt types for systematic agent security evaluation.",
        "structure": "6 attack prompt types: combined_attack, context_ignoring, escape_characters, fake_completion, naive, plus average. Measures Utility (%) and ASR (%). Used by both Progent and ToolSafe for evaluation.",
        "paper_usage": "Used as secondary benchmark by Progent (evaluated against 3 defenses) and as primary benchmark by ToolSafe for evaluating TS-Guard's step-level proactive guardrail.",
        "project_value": "The 6 attack prompt types provide a classification scheme for input-level risk assessment in the MCP security system. The utility vs ASR tradeoff measurements help calibrate how much task degradation is acceptable for security enforcement.",
    },
    {
        "dataset_name": "AgentHarm Dataset",
        "paper_title": "AgentHarm: Harmful Agent Behavior Benchmark (Gray Swan / UK AISI, 2025)",
        "link": "https://huggingface.co/datasets/ai-safety-institute/AgentHarm",
        "description": "Benchmark for evaluating harmful agent behaviors including deepfakes, misinformation, and other dangerous capabilities. Contains ReAct trajectories with harmful agent requests.",
        "structure": "ReAct-format trajectories with harmful behavior requests. Covers multiple harm categories. Used by ToolSafe and TraceAegis for evaluation.",
        "paper_usage": "Used as evaluation benchmark by ToolSafe (measuring ASR, Safety Rate, and Helpfulness) and referenced by TraceAegis as a related benchmark for harmful behavior detection.",
        "project_value": "Provides ground truth for the most severe risk tier — requests that should always be denied. The harm categories inform the risk scorer's highest-severity classification criteria.",
    },
    {
        "dataset_name": "RAS-Eval",
        "paper_title": "RAS-Eval: Comprehensive Benchmark for LLM Agent Security (Fu et al., 2025)",
        "link": "https://github.com/lanzer-tree/RAS-Eval",
        "description": "Comprehensive benchmark for security evaluation of LLM agents in real-world environments. Contains CFI (Control-Flow Integrity) violation attacks adapted for metadata poisoning evaluation.",
        "structure": "Multi-category security evaluation. Per-model distributions vary (e.g., Qwen3-8b: 919 Normal, 6 Poisoned). Format adapted for metadata poisoning by MindGuard evaluation.",
        "paper_usage": "Used by MindGuard as a third evaluation dataset with attack payloads transformed into metadata poisoning format. Also referenced by MCP-Guard as an evaluation corpus.",
        "project_value": "The control-flow integrity violation focus is relevant to detecting tool invocation sequence manipulation in MCP. Provides additional evaluation data for the risk scorer's ability to detect metadata-level attacks.",
    },
    {
        "dataset_name": "EICU-AC (Healthcare Access Control)",
        "paper_title": "GuardAgent: Safeguard LLM Agents by a Guard Agent via Knowledge-Enabled Reasoning (Xiang et al., 2024)",
        "link": "https://github.com/guardagent/code",
        "description": "Access control benchmark for healthcare LLM agents built on the eICU Collaborative Research Database. Defines three roles (physician, nursing, general admin) with different database permissions for role-based access evaluation.",
        "structure": "100 test samples across 10 databases. Three roles with distinct permission levels. Based on the eICU clinical database.",
        "paper_usage": "Primary benchmark for evaluating GuardAgent's access control enforcement in the EHRAgent healthcare scenario. GuardAgent achieved >98% accuracy in safety assessment.",
        "project_value": "Demonstrates role-based access control evaluation methodology directly applicable to MCP agent access control. The three-role permission model provides a template for MCP agent role classification. Healthcare domain demonstrates high-stakes access control requirements.",
    },
    {
        "dataset_name": "Mind2Web-SC (Web Navigation Safety Control)",
        "paper_title": "GuardAgent: Safeguard LLM Agents by a Guard Agent via Knowledge-Enabled Reasoning (Xiang et al., 2024)",
        "link": "https://github.com/guardagent/code",
        "description": "Safety control benchmark for web navigation agents built on Mind2Web. Defines safety rules based on user attributes (age, license) and action categories (shopping, rentals, etc.) for web interaction safety evaluation.",
        "structure": "100 test samples across 7 action categories. Safety rules based on user attributes and action types. Built on the Mind2Web web navigation dataset.",
        "paper_usage": "Primary benchmark for evaluating GuardAgent's safety control in the SeeAct web navigation scenario. Tests whether the guard agent correctly enforces attribute-based safety constraints.",
        "project_value": "The attribute-based safety rules provide a model for context-dependent risk scoring in MCP — where the same tool invocation may be safe or risky depending on the requesting agent's attributes and context.",
    },
    {
        "dataset_name": "MiniScope Synthetic Permission Dataset",
        "paper_title": "MiniScope: A Least Privilege Framework for Authorizing Tool Calling Agents (Zhu et al., 2025)",
        "link": "https://arxiv.org/abs/2512.11147",
        "description": "Synthetic dataset of realistic user-agent interactions derived from 10 popular real-world applications. Contains single-method and multi-method requests for evaluating permission minimality, runtime overhead, and user effort.",
        "structure": "10 applications (Gmail: 79 methods/10 scopes, Google Calendar: 37/17, Slack: 247/84, Dropbox: 120/13, etc.). Single-method requests per API method. 200 multi-method requests per app. 2 multi-app suites (171 and 465 methods). 1-6% latency overhead measured.",
        "paper_usage": "Evaluated MiniScope against 3 baselines (Vanilla, PerMethod, LLMScope) on permission minimality, runtime overhead, and user confirmation rates. MiniScope achieved lowest mismatch rates.",
        "project_value": "The permission hierarchy data from 10 real applications provides reference architectures for MCP tool permission modeling. The scope-based access control approach directly maps to MCP capability-level risk assessment.",
    },
    {
        "dataset_name": "Trust Paradox Evaluation Scenarios (19 scenarios)",
        "paper_title": "The Trust Paradox in LLM-Based Multi-Agent Systems (Xu et al., 2025)",
        "link": "https://arxiv.org/abs/2510.18563",
        "description": "Carefully constructed evaluation scenarios validating the trust paradox — that granting agents more capability inherently increases their vulnerability surface. Spans 5 capability levels tested across 4 model backends.",
        "structure": "19 scenarios across 5 capability tiers (basic, intermediate, advanced, expert, critical). 4 LLM backends: GPT-4o, Claude 3.5 Sonnet, Llama 3.1 70B, Mixtral 8x22B. Metrics: Trust Calibration Index, Capability-Risk Ratio, Safety Violation Rate.",
        "paper_usage": "Empirically validated the trust paradox with TCI ranging 0.72-0.89, confirming systematic trust miscalibration. Inter-agent trust scores stabilize at 0.73-0.91 but require 8-15 iterations.",
        "project_value": "Demonstrates that the risk scorer must account for the trust-capability paradox — more capable agents may need higher scrutiny, not less. The 5-tier capability model informs how the risk scorer should adjust severity based on agent capability level.",
    },
    {
        "dataset_name": "Indirect PI Attack Dataset (1,068 instances)",
        "paper_title": "Exploiting Web Search Tools of AI Agents for Data Exfiltration (Rall et al., 2025)",
        "link": "https://anonymous.4open.science/r/web-search-exploit-paper-FFC6",
        "description": "Systematic evaluation dataset for indirect prompt injection via web search tools in AI agents. Generated from 89 attack templates with 12 variations each using PyRIT converters (LLM fuzzing, encoding, Unicode obfuscation, etc.).",
        "structure": "1,068 unique attack instances from 89 templates x 12 variations. Tested on 28 models across providers (X-AI, Inception, Qwen, Meta, Anthropic, OpenAI, Google, DeepSeek, Mistral). Variations: LLM-based fuzzing, Base64 encoding, Unicode obfuscation, random capitalization, translation.",
        "paper_usage": "Systematically evaluated LLM vulnerability to data exfiltration via web search tools. Found X-AI grok-4 most vulnerable (72.4%), OpenAI/Google/Amazon most resilient (near 0%).",
        "project_value": "The 12 attack variation types provide obfuscation categories the risk scorer must handle. The per-model vulnerability data informs model-specific risk adjustment. The PyRIT-based generation approach can be adapted to generate MCP-specific attack variants for training data.",
    },
    {
        "dataset_name": "Synthetic PI Dataset (500 prompts)",
        "paper_title": "Prompt Injection Detection and Mitigation via AI Multi-Agent NLP Frameworks (Gosmar et al., 2025)",
        "link": "https://arxiv.org/abs/2503.11517",
        "description": "500 engineered injection prompts across 10 attack categories generated by OpenAI o3-mini. Categories include Direct Override, Authority Assertions, Hidden Commands, Role-Play, Logical Traps, Multi-Step, Conflicting Instructions, HTML/Markdown Embeds, Hybrid, and Social Engineering.",
        "structure": "500 prompts total (50 per category). 10 attack categories. Processed through a 3-agent pipeline (Generator, Sanitizer, Enforcer). Novel metrics: ISR, POF, PSR, CCS, TIVS.",
        "paper_usage": "Evaluated the multi-agent NLP defense framework achieving 45.7% reduction in vulnerability score through layered mitigation. TIVS composite score measured across all categories.",
        "project_value": "The 10 attack categories provide a prompt-level risk classification taxonomy. The TIVS composite scoring methodology is directly transferable to MCP risk scoring design. The 4 novel metrics (ISR, POF, PSR, CCS) can be adapted as sub-scores in the MCP risk model.",
    },
    {
        "dataset_name": "ShareGPT Conversation Dataset",
        "paper_title": "Imprompter: Tricking LLM Agents into Improper Tool Use (Fu et al., 2024)",
        "link": "https://huggingface.co/datasets/anon8231489123/ShareGPT_Vicuna_unfiltered",
        "description": "Real human-ChatGPT 3.5 conversations from late 2022 to early 2023. Used as training data for information exfiltration attacks that trick agents into improper tool use.",
        "structure": "~53,000 conversations. First 100 sampled for training, last 100 for validation. Format: multi-turn conversation transcripts.",
        "paper_usage": "Used to train D_text for information exfiltration attacks. Evaluated using Syntax Correctness, PPL, Word Extraction Precision, and Word Extraction GPT Score.",
        "project_value": "Provides realistic conversation patterns for testing whether the risk scorer can detect exfiltration attempts embedded in natural-looking conversations. The adversarial attack methodology demonstrates attacks the risk scorer must handle.",
    },
    {
        "dataset_name": "WildChat Dataset (1M interactions)",
        "paper_title": "Imprompter: Tricking LLM Agents into Improper Tool Use (Fu et al., 2024)",
        "link": "https://huggingface.co/datasets/allenai/WildChat-1M",
        "description": "1 million real human-LLM interaction logs from 2024 with PII annotations. Used for training and validating PII exfiltration attacks. Full version with toxic data requires gated access.",
        "structure": "1M interaction logs. 49 conversations with PII found (24 train, 25 validation). PII annotations included. Gated access for toxic content subset.",
        "paper_usage": "Used for training/validation of PII exfiltration attacks. Evaluated with Syntax Correctness, PII Precision, PII Recall, and Context GPT Score.",
        "project_value": "The PII-annotated data is useful for testing whether the risk scorer can detect privacy leakage risks in MCP tool invocations. The large scale (1M interactions) provides diverse patterns for robust evaluation.",
    },
    # ── Trustworthiness & General Datasets ──
    {
        "dataset_name": "DecodingTrust Dataset",
        "paper_title": "DecodingTrust: A Comprehensive Assessment of Trustworthiness in GPT Models (Wang et al., 2023)",
        "link": "https://huggingface.co/datasets/AI-Secure/DecodingTrust",
        "description": "Eight-dimensional trustworthiness evaluation dataset for GPT models covering toxicity, stereotype bias, adversarial robustness, OOD robustness, adversarial demonstrations, privacy, machine ethics, and fairness. NeurIPS 2023 Outstanding Paper.",
        "structure": "8 evaluation dimensions, each with multiple sub-datasets. Includes: RealToxicityPrompts, BOLD, BBQ, SST-2, AdvGLUE, ANLI, RealtimeQA, MMLU, Enron Email, ETHICS, Jiminy Cricket, Adult (UCI), and more. Available on HuggingFace.",
        "paper_usage": "Evaluated GPT-4 and GPT-3.5 across all 8 dimensions. Discovered GPT-4 is more vulnerable to jailbreaking than GPT-3.5 despite better baseline performance.",
        "project_value": "The 8-dimension framework provides a ready-made scoring rubric adaptable for MCP agent trustworthiness. The finding that capability does not equal trustworthiness is critical for risk assessment — the risk scorer should not assume more capable models are safer.",
    },
    {
        "dataset_name": "TrustLLM Benchmark Datasets (30+)",
        "paper_title": "TrustLLM: Trustworthiness in Large Language Models (Huang et al., 2024)",
        "link": "https://github.com/HowieHwong/TrustLLM",
        "description": "Comprehensive trustworthiness benchmark evaluating 16 LLMs across six dimensions (truthfulness, safety, fairness, robustness, privacy, machine ethics) using over 30 datasets. Published at ICML 2024.",
        "structure": "30+ datasets across 6 dimensions and 18 subcategories. Includes: SQuAD2.0, TruthfulQA, HaluEval, CrowS-Pair, StereoSet, AdvGLUE, Enron Email, ETHICS, ConfAIde, MoralChoice, Social Chemistry 101, and many more. 16 LLMs evaluated.",
        "paper_usage": "Largest-scale LLM trustworthiness study. Found proprietary LLMs generally more trustworthy but some over-optimize safety at the cost of utility.",
        "project_value": "The 6 trustworthiness dimensions map to MCP risk scoring dimensions. The 30+ datasets provide extensive evaluation data. The safety-utility tradeoff finding is directly relevant to calibrating MCP risk scoring thresholds.",
    },
    {
        "dataset_name": "NVD/CVE Database (31,000+ entries)",
        "paper_title": "From Description to Score: Can LLMs Quantify Vulnerabilities? (Jafarikhah et al., 2026)",
        "link": "https://nvd.nist.gov/",
        "description": "National Vulnerability Database containing CVE entries with CVSS v3.1 base metrics. Used to evaluate whether LLMs can automatically score vulnerabilities from textual descriptions.",
        "structure": "31,000+ CVEs from 2019-2024 with CVSS v3.1 base metrics (8 sub-metrics). Tested with 6 LLMs: GPT-4o, GPT-5, Llama-3.3-70B, Gemini-2.5-Flash, DeepSeek-R1, Grok-3. Format: CVE description text paired with CVSS scores.",
        "paper_usage": "Evaluated 6 LLMs for automated CVSS scoring. Found two-shot prompting optimal. GPT-5 achieved highest precision. Meta-classifier combining LLM outputs provided marginal improvement.",
        "project_value": "CVSS scoring methodology is directly transferable to MCP risk scoring. The approach of using LLMs to quantify severity from textual descriptions mirrors how MCP Security could score agent requests from tool descriptions. The ensemble meta-classifier approach applies to combining multiple risk signals.",
    },
    {
        "dataset_name": "glaive-function-calling-v2",
        "paper_title": "MCIP: Protecting MCP Safety via Model Contextual Integrity Protocol (Jing et al., 2025)",
        "link": "https://huggingface.co/datasets/glaiveai/glaive-function-calling-v2",
        "description": "Widely used open-source dataset for training models on function calling. Contains real function calling conversations released by Glaive AI.",
        "structure": "112,960 instances total. Multi-turn function calling conversations. Used as source for both MCIP-Bench construction (200 gold instances) and MCIP Guardian training (2,000 sampled rows). Function pool of 10,633 function call pairs.",
        "paper_usage": "Primary source for constructing MCIP-Bench gold instances and MCIP Guardian training data. Provided realistic function calling patterns for safety evaluation.",
        "project_value": "Large-scale function calling data directly relevant to training an MCP tool invocation risk classifier. The 112K+ instances provide diverse calling patterns for learning normal behavior baselines. The function pool enables training risk models on realistic tool interaction sequences.",
    },
    {
        "dataset_name": "ToolACE Dataset",
        "paper_title": "Used by MCIP (Jing et al., 2025) and MindGuard (Wang et al., 2025)",
        "link": "https://huggingface.co/datasets/Team-ACE/ToolACE",
        "description": "Function calling dataset used as a complementary data source for MCP security evaluation. Provides tool call samples for generalization validation and clean baseline comparison.",
        "structure": "11,300 rows. Used by MCIP-Bench for supplementary evaluation (1,026 instances). Also used by MindGuard as part of a 12,000+ heterogeneous clean tool call validation set.",
        "paper_usage": "Used by MCIP for generalization validation on unseen risks. Used by MindGuard for TAE (Total Attention Energy) signal validation, confirming the signal can distinguish decision sources from non-sources.",
        "project_value": "Provides clean/benign tool call baselines essential for training a risk scorer to distinguish normal from anomalous MCP tool invocations. The generalization validation demonstrates how to test risk scoring robustness across different data sources.",
    },
    {
        "dataset_name": "ToolBench Dataset",
        "paper_title": "ToolHijacker: Prompt Injection Attack to Tool Selection in LLM Agents (Shi et al., 2025)",
        "link": "https://github.com/OpenBMB/ToolBench",
        "description": "Large-scale benchmark for enhancing tool-use capabilities of LLMs with instruction-tuning samples. Contains tools from RapidAPI used for evaluating tool selection attacks and defenses.",
        "structure": "126,486 instruction-tuning samples leveraging 16,464 tool documents from RapidAPI. After deduplication: 9,650 benign tool documents. Used with 10 target tasks and 1,000 task descriptions for attack evaluation.",
        "paper_usage": "Secondary evaluation dataset for ToolHijacker. Gradient-free attack achieved 88.2% ASR with GPT-4o. Larger tool library (9,650 docs) made defense evaluation more realistic. Also used by MindGuard for TAE validation.",
        "project_value": "The 16K+ tool documents provide a large-scale tool registry for testing MCP tool-level risk classification. The ToolHijacker results (88.2% ASR) quantify the risk of tool selection manipulation that the MCP risk scorer must detect.",
    },
    {
        "dataset_name": "MetaTool Benchmark Dataset",
        "paper_title": "ToolHijacker: Prompt Injection Attack to Tool Selection in LLM Agents (Shi et al., 2025)",
        "link": "https://github.com/HowieHwong/MetaTool",
        "description": "Benchmark focusing on LLMs' capabilities in tool usage with scenario-driven evaluations. Contains tool documents sourced from OpenAI Plugins for evaluating tool selection accuracy and attack vulnerability.",
        "structure": "21,127 instances involving 199 benign tool documents from OpenAI Plugins. 10 high-quality target tasks designed for real-world needs. 100 target task descriptions per task (1,000 total). Metrics: Accuracy, ASR, Hit Rate, Attack Hit Rate.",
        "paper_usage": "Primary evaluation dataset for ToolHijacker. Gradient-free attack achieved 96.7% ASR with GPT-4o on MetaTool. Human study with 6 participants showed >=71% failure to detect malicious tool documents.",
        "project_value": "The 199 tool document library from OpenAI Plugins provides realistic tool metadata for training the MCP risk scorer. The 96.7% ASR demonstrates that tool selection attacks are highly effective — the risk scorer must prioritize tool integrity verification.",
    },
    {
        "dataset_name": "Meta Tool-Use Agentic PI Benchmark (600 scenarios)",
        "paper_title": "LlamaFirewall: An Open Source Guardrail System for Building Secure AI Agents (Chennabasappa et al., 2025)",
        "link": "https://huggingface.co/datasets/facebook/llamafirewall-alignmentcheck-evals",
        "description": "Novel benchmark for evaluating prompt injection resilience in agentic environments across travel, information retrieval, and productivity domains. Simulates smartphone apps with realistic tool-use scenarios.",
        "structure": "600 scenarios (300 benign, 300 malicious). 7 injection techniques. 8 threat categories. Three domains: travel, info-retrieval, productivity. Format: tool-use agentic interaction traces.",
        "paper_usage": "Used to evaluate LlamaFirewall's combined PromptGuard + AlignmentCheck pipeline for PI resilience. Measured ASR per attack type and threat category.",
        "project_value": "The 600-scenario benchmark with balanced benign/malicious split provides well-structured evaluation data for the MCP risk scorer. The 7 injection techniques and 8 threat categories inform the risk scoring taxonomy. Available on HuggingFace for direct use.",
    },
    {
        "dataset_name": "CyberSecEval3",
        "paper_title": "LlamaFirewall: An Open Source Guardrail System for Building Secure AI Agents (Chennabasappa et al., 2025)",
        "link": "https://github.com/meta-llama/PurpleLlama/tree/main/CybersecurityBenchmarks",
        "description": "Cybersecurity evaluation benchmark for LLM-generated code from Meta's PurpleLlama project. Tests code completion safety across multiple programming languages.",
        "structure": "50 code completions per language across 7+ programming languages. Evaluates whether LLM-generated code contains vulnerabilities. Metrics: Precision (96%), Recall (79%).",
        "paper_usage": "Used to evaluate CodeShield component of LlamaFirewall for detecting insecure code patterns in LLM-generated code.",
        "project_value": "Code security evaluation is relevant when MCP agents generate or modify code through tool invocations. The risk scorer could integrate similar static analysis to assess code-generating tool calls.",
    },
    {
        "dataset_name": "Anthropic Red-Teaming Dataset",
        "paper_title": "NeMo Guardrails: A Toolkit for Controllable and Safe LLM Applications (Rebedea et al., 2023)",
        "link": "https://github.com/anthropics/hh-rlhf",
        "description": "Human-annotated prompts rated 0-4 for alignment bypass attempts, used to evaluate moderation rails. Paired with the Anthropic Helpful Dataset for balanced safety evaluation.",
        "structure": "Highest-harm prompts sampled from the full red-teaming set. Balanced with ~200 total samples (harmful + helpful). Format: text prompts with harm severity ratings.",
        "paper_usage": "Used to evaluate NeMo Guardrails' moderation capability — measuring percentage of harmful content blocked versus helpful content correctly allowed.",
        "project_value": "Provides adversarial prompts for testing the MCP risk scorer's ability to detect harmful requests. The harm severity ratings (0-4) provide labeled data for training graduated risk classification.",
    },
    {
        "dataset_name": "CIAQA (Compositional Instruction Attack QA)",
        "paper_title": "From Prompt Injections to Protocol Exploits: Threats in LLM-Powered AI Agents Workflows (Ferrag et al., 2025)",
        "link": "https://arxiv.org/abs/2506.23260",
        "description": "Dataset for evaluating Compositional Instruction Attacks (CIA) on LLMs, containing multiple-choice questions based on successful jailbreaks. Demonstrates >95% ASR on major models.",
        "structure": "2,700 multiple-choice questions based on 900 successful jailbreaks. Format: QA pairs testing compositional attack effectiveness.",
        "paper_usage": "Cited to demonstrate that CIA achieves >95% ASR on GPT-4, GPT-3.5, and Llama2-70b-chat, quantifying the severity of compositional instruction attacks.",
        "project_value": "Compositional attacks that achieve >95% ASR represent a critical threat vector for MCP security. The risk scorer must detect when multiple innocent-looking instructions combine into an attack pattern.",
    },
    {
        "dataset_name": "CrAIBench (Web3 AI Agent Benchmark)",
        "paper_title": "From Prompt Injections to Protocol Exploits: Threats in LLM-Powered AI Agents Workflows (Ferrag et al., 2025)",
        "link": "https://arxiv.org/abs/2506.23260",
        "description": "Web3 domain-specific benchmark for assessing AI agent robustness against context manipulation attacks across realistic blockchain tasks including token transfers, trading, and cross-chain interactions.",
        "structure": "150+ realistic blockchain tasks. 500+ attack test cases. Covers token transfers, trading, cross-chain interactions. Tests both prompt-based and fine-tuning-based defenses.",
        "paper_usage": "Cited to demonstrate context manipulation attack effectiveness in Web3 ecosystems. Found prompt-based defenses ineffective while fine-tuning-based defenses showed more robustness.",
        "project_value": "Relevant to MCP risk scoring in financial/blockchain tool contexts. Demonstrates that domain-specific attack testing is necessary — the risk scorer may need specialized assessment for high-value MCP tool categories.",
    },
    {
        "dataset_name": "AlpacaFarm Dataset (208 samples)",
        "paper_title": "Defense Against Prompt Injection Attack by Leveraging Attack Techniques (Chen et al., 2025)",
        "link": "https://github.com/tatsu-lab/stanford_alpaca",
        "description": "Instruction-following dataset subset used for evaluating direct prompt injection defense. The defense approach inverts attack techniques for defensive purposes without requiring model fine-tuning.",
        "structure": "208 samples used for direct prompt injection evaluation. Part of the larger Stanford Alpaca dataset. Paired with 2,000-sample Filtered QA Dataset for indirect PI evaluation.",
        "paper_usage": "Used to evaluate training-free defense achieving state-of-the-art results by inverting prompt injection attack techniques for defensive purposes.",
        "project_value": "The training-free defense approach is highly relevant for MCP Security since it requires no model modification. The inverted-attack methodology could be adapted for real-time MCP request screening.",
    },
    {
        "dataset_name": "BFCL-v3 (Berkeley Function Calling Leaderboard)",
        "paper_title": "MCIP: Protecting MCP Safety via Model Contextual Integrity Protocol (Jing et al., 2025)",
        "link": "https://huggingface.co/datasets/gorilla-llm/Berkeley-Function-Calling-Leaderboard",
        "description": "Benchmark for evaluating function calling capabilities of LLMs. Used to assess whether safety-oriented MCP designs affect practical function calling utility.",
        "structure": "Comprehensive function calling evaluation across multiple LLM models. Measures overall accuracy percentage. Available on both GitHub and HuggingFace.",
        "paper_usage": "Used as utility metric to evaluate the safety-utility tradeoff of MCIP Guardian — ensuring safety improvements don't degrade function calling capability.",
        "project_value": "Essential for measuring whether MCP risk scoring introduces unacceptable utility degradation. A risk scorer that blocks too many legitimate calls is impractical. BFCL-v3 provides the utility measurement standard.",
    },
]

# ---------------------------------------------------------------------------
# SECTION 2: BENCHMARKS DATA
# ---------------------------------------------------------------------------

BENCHMARKS: list[dict[str, str]] = [
    # ── MCP-Specific Benchmarks ──
    {
        "benchmark_name": "MCPSecBench",
        "paper_title": "MCPSecBench: A Systematic Security Benchmark and Playground for Testing Model Context Protocols (Yang et al., 2025)",
        "link": "https://arxiv.org/abs/2508.13220",
        "description": "The first systematic MCP security benchmark providing taxonomy-driven test cases, GUI test harness, prompt datasets, attack scripts, and server/client implementations for reproducible evaluation across 17 attack types and 4 attack surfaces.",
        "metrics": "Attack Success Rate (ASR), Refusal Rate (RR), cost per testing round ($0.41-$0.76). Evaluated per attack type per platform.",
        "baselines": "Claude Desktop, OpenAI GPT-4.1, Cursor v2.3.29. Protocol-side: 100% universal ASR. Client-side: ~33% average. Server-side: ~47%. Host-side: ~27%.",
        "how_to_use": "Clone the benchmark repository. Configure MCP client/server test environments. Run attack scripts against target platforms. Use the GUI test harness for automated benchmarking. Requires access to target LLM APIs and MCP client installations. 15 trials per attack for statistical robustness.",
        "project_value": "The definitive benchmark for evaluating MCP security solutions. The 4-surface taxonomy (client/protocol/server/host) maps directly to the risk scoring system's evaluation dimensions. Essential for comparing the MCP risk scorer against published baselines.",
    },
    {
        "benchmark_name": "MCPTox",
        "paper_title": "MCPTox: A Benchmark for Tool Poisoning Attack on Real-World MCP Servers (Wang et al., 2025)",
        "link": "https://openreview.net/forum?id=xbs5dVGUQ8",
        "description": "First benchmark for systematic evaluation of tool poisoning attacks on real-world MCP servers. Tests three attack paradigms (Explicit Trigger, Implicit Trigger — Function Hijacking, Implicit Trigger — Parameter Tampering) across 10 attack categories using authentic MCP server tools.",
        "metrics": "Attack Success Rate (ASR), Refusal Rate (RR). Per-paradigm and per-model breakdowns. Overall 72.8% ASR measured.",
        "baselines": "GPT-4o (61.8% ASR), Qwen3-32b (58.5%), Claude (34.3%). Reasoning mode impact: Qwen3-8b shows +27.8% ASR with reasoning enabled.",
        "how_to_use": "Access the benchmark dataset with 45 real-world MCP servers and 353 tools. Generate malicious test cases using the three attack templates. Run against target LLM agents. Measure ASR and RR per model. Requires API access to target LLMs and MCP server configurations.",
        "project_value": "Essential for testing the risk scorer's ability to detect tool poisoning — the most direct MCP attack vector. The three paradigms test different detection difficulty levels. Real-world server data ensures practical relevance of evaluation results.",
    },
    {
        "benchmark_name": "MCP-SafetyBench",
        "paper_title": "MCP-SafetyBench: A Benchmark for Safety Evaluation of Large Language Models with Real-World MCP Servers (Zong et al., 2025)",
        "link": "https://arxiv.org/abs/2512.15163",
        "description": "Multi-domain MCP safety benchmark evaluating 13 LLMs across 5 real-world domains (Location Navigation, Repository Management, Financial Analysis, Browser Automation, Web Search) with 20 attack types in multi-turn interactions.",
        "metrics": "Task Success Rate (TSR), Attack Success Rate (ASR), Defense Success Rate (DSR), safety prompt effectiveness.",
        "baselines": "13 LLMs: GPT-5, GPT-4.1, GPT-4o, o4-mini, Claude-3.7-Sonnet, Claude-4.0-Sonnet, Gemini-2.5-Pro, Gemini-2.5-Flash, Grok-4, GLM-4.5, Kimi-K2, Qwen3-235B, DeepSeek-V3.1.",
        "how_to_use": "Configure real-world MCP servers for 5 domains. Apply adversarial tool description modifications. Run multi-turn evaluation sessions with target LLMs. Measure TSR and ASR per domain-attack combination. Requires API access to all 13 models and domain-specific MCP server setups.",
        "project_value": "The broadest MCP safety benchmark by domain coverage. The 5-domain evaluation ensures the risk scorer generalizes across MCP use cases. Multi-turn interaction testing evaluates sustained risk assessment accuracy over extended sessions.",
    },
    {
        "benchmark_name": "MCP-AttackBench",
        "paper_title": "MCP-Guard: A Multi-Stage Defense-in-Depth Framework for Securing MCP in Agentic AI (Xing et al., 2025)",
        "link": "https://arxiv.org/abs/2508.10991",
        "description": "Large-scale attack benchmark with hierarchical taxonomy of MCP threats organized into Semantic & Adversarial, Protocol-Specific, and Injection & Execution families. Designed for training and evaluating cascaded defense systems.",
        "metrics": "F1-score (95.4% achieved), Attack Success Rate, Detection Success Rate, runtime latency (455.9ms avg, 50.9ms optimized), speedup ratio (2.04x vs standalone LLM).",
        "baselines": "Standalone LLM defenses. Three-stage cascade: Stage I pattern matching (38.9% filter rate, <2ms), Stage II neural detection (96.01% F1), Stage III LLM arbitration.",
        "how_to_use": "Use the 70,448-sample dataset for training detection models. Implement the three-stage cascade pipeline. Fine-tune E5 embedding model on training corpus (5,258 samples). Evaluate with binary cross-entropy loss. Requires GPU for neural model training and LLM API access for Stage III.",
        "project_value": "Provides both training data and a defense architecture blueprint for the MCP risk scoring system. The cascaded approach (fast pattern → neural → LLM) directly maps to a multi-stage risk scoring pipeline. The 95.4% F1 sets the accuracy target.",
    },
    {
        "benchmark_name": "MCIP-Bench",
        "paper_title": "MCIP: Protecting MCP Safety via Model Contextual Integrity Protocol (Jing et al., 2025)",
        "link": "https://github.com/HKUST-KnowComp/MCIP",
        "description": "MCP contextual integrity benchmark evaluating LLMs' ability to detect 10 risk types in MCP function-calling interactions. Tests both Safety Awareness (binary safe/unsafe) and Risk Resistance (11-class identification).",
        "metrics": "Safety Awareness Accuracy, Safety Awareness Macro-F1, Risk Resistance Accuracy, Risk Resistance Macro-F1, BFCL Utility (overall accuracy %).",
        "baselines": "xLAM-2-70b-fc-r, xLAM-2-32b-fc-r, xLAM-2-8B-fc-r, ToolACE-2-8B, Qwen2.5-72B-Instruct, Qwen2.5-32B-Instruct, DeepSeek-R1.",
        "how_to_use": "Clone the MCIP GitHub repository. Load the 2,218-instance benchmark. Run target LLMs on both Safety Awareness and Risk Resistance tasks. Compare against published baseline results. Requires API access to evaluation models and local compute for open-source models.",
        "project_value": "The 10 risk types provide a granular classification scheme for MCP risk scoring. The dual evaluation (binary + 11-class) tests both coarse and fine-grained risk detection. Available on GitHub for immediate use.",
    },
    {
        "benchmark_name": "ProtoAMP (Protocol Amplification Benchmark)",
        "paper_title": "Breaking the Protocol: Security Analysis of the MCP Specification and Prompt Injection Vulnerabilities (Maloyan & Namiot, 2026)",
        "link": "https://arxiv.org/abs/2601.17549",
        "description": "Benchmark quantifying how MCP architecture amplifies prompt injection attack success rates compared to non-MCP baselines. Tests attacks at three protocol layers (resource content, tool response payloads, sampling request prompts) with the proposed AttestMCP defense.",
        "metrics": "Attack Success Rate (ASR), amplification factor (% increase vs baseline), defense reduction (ASR before/after AttestMCP), latency overhead (8.3ms cold, 2.4ms warm).",
        "baselines": "Non-MCP baseline ASR. Three MCP server implementations: filesystem, git, sqlite. Four LLM backends: Claude-3.5-Sonnet, GPT-4o, Llama-3.1-70B. AttestMCP defense reduces ASR from 52.8% to 12.4% (76.5% reduction).",
        "how_to_use": "Set up 3 MCP server implementations. Configure 4 LLM backends. Run 847 test scenarios measuring ASR at each protocol layer. Compare MCP vs non-MCP baselines. Implement AttestMCP extension for defense evaluation. Requires MCP server infrastructure and LLM API access.",
        "project_value": "Quantifies the MCP-specific risk amplification that the risk scorer must account for. The +23-41% amplification factor demonstrates that protocol context should increase risk scores. AttestMCP's 76.5% ASR reduction shows the value of protocol-level defenses.",
    },
    {
        "benchmark_name": "AttestMCP",
        "paper_title": "Breaking the Protocol: Security Analysis of the MCP Specification and Prompt Injection Vulnerabilities (Maloyan & Namiot, 2026)",
        "link": "https://arxiv.org/abs/2601.17549",
        "description": "Backward-compatible attestation extension for MCP providing 5 security enhancements: Capability Attestation (cryptographic proof), Message Authentication (HMAC-SHA256), Origin Tagging (server identification), Isolation Enforcement (cross-server data flow restrictions), and Replay Protection (nonce-based time windows).",
        "metrics": "ASR reduction (52.8% → 12.4%, 76.5% reduction), performance overhead (8.3ms cold start, 2.4ms warm cache per message).",
        "baselines": "Unprotected MCP protocol baseline. Compared against: indirect injection (47.8% → reduced), cross-server propagation (61.3% → reduced), sampling vulnerability (67.2% → reduced).",
        "how_to_use": "Implement the 5 AttestMCP extensions on MCP servers. Add HMAC-SHA256 message signing. Configure nonce-based replay protection. Deploy origin tags on all server responses. Requires cryptographic key management infrastructure and protocol-level modifications.",
        "project_value": "AttestMCP's 5 security features inform what the risk scorer should verify: capability attestation, message integrity, origin tracking, isolation, and replay protection. Each missing feature could increase risk scores.",
    },
    {
        "benchmark_name": "MCLIB (MCP Attack Library)",
        "paper_title": "Referenced by When MCP Servers Attack (Zhao et al., 2025); Guo et al., 2025",
        "link": "https://arxiv.org/abs/2508.12538",
        "description": "Unified plugin-based framework for simulating and evaluating MCP attacks. Catalogs 31 attacks across four families: direct tool injection, indirect tool injection, malicious-user attacks, and LLM-inherent attacks.",
        "metrics": "Attack success rates per family and per attack type. Plugin-based evaluation metrics.",
        "baselines": "31 attacks across 4 families provide comprehensive coverage of the MCP threat landscape.",
        "how_to_use": "Install the MCLIB framework. Configure attack plugins for target MCP environments. Run systematic evaluation across all 4 attack families. Requires MCP server/client setup and LLM access.",
        "project_value": "The 31-attack catalog across 4 families provides the most comprehensive threat enumeration for the risk scoring taxonomy. The plugin architecture enables systematic testing of risk scorer coverage against all known attack types.",
    },
    {
        "benchmark_name": "MCPSafetyScanner",
        "paper_title": "MCP Safety Audit: LLMs with the MCP Allow Major Security Exploits (Radosevich & Halloran, 2025)",
        "link": "https://github.com/johnhalloran321/mcpSafetyScanner",
        "description": "First automated MCP security auditing tool. Scans MCP deployments for vulnerability patterns and executes proof-of-concept exploits including credential theft, malicious code execution, and retrieval-agent deception attacks.",
        "metrics": "Qualitative exploit validation (credential theft success, code execution success). Tested on Claude 3.7 and Llama-3.3-70B-Instruct.",
        "baselines": "Tested against standard MCP servers: filesystem, Slack, Everything, Chroma. Attack types: MCE, RAC, CT, RADE.",
        "how_to_use": "Install MCPSafetyScanner from GitHub. Point it at target MCP server deployments. Run automated vulnerability scanning. Review exploit reports. Requires Node.js and access to target MCP servers.",
        "project_value": "A ready-made validation tool for the risk scoring system — run MCPSafetyScanner against test deployments and verify the risk scorer flags the same vulnerabilities. The documented attack types define critical severity categories.",
    },
    # ── Agent Security Benchmarks ──
    {
        "benchmark_name": "AgentDojo",
        "paper_title": "AgentDojo: A Dynamic Environment to Evaluate Prompt Injection Attacks and Defenses for LLM Agents (Debenedetti et al., 2024)",
        "link": "https://github.com/ethz-spylab/agentdojo",
        "description": "The most widely adopted dynamic evaluation environment for prompt injection attacks and defenses in tool-using LLM agents. Provides 4 realistic task suites with parameterized attacks, defense baselines, and reproducible benchmarking. Used by Progent, LlamaFirewall, TraceAegis, ToolSafe, and The Task Shield.",
        "metrics": "Utility (task completion %), Utility under attack, Attack Success Rate (ASR), AUC, Recall@1%FPR, TPR, FPR, Balanced Accuracy, F1.",
        "baselines": "Multiple defenses compared: repeat_user_prompt, spotlighting_with_delimiting, tool_filter, transformers_pi_detector, DataSentinel, Llama Prompt Guard 2. Progent achieved ASR reduction from 41.2% to 2.2%.",
        "how_to_use": "pip install agentdojo. Implement agent within the framework. Run against 4 task suites (Workspace, Travel, Banking, Slack) with 97 user tasks and 629 injection tasks. Compare utility vs ASR tradeoff. Requires Python 3.10+ and LLM API access.",
        "project_value": "The standard benchmark for agent security evaluation — must be included for credibility. The 4-domain coverage tests generalization. The extensive defense baseline comparisons enable positioning the MCP risk scorer against state-of-the-art.",
    },
    {
        "benchmark_name": "InjecAgent",
        "paper_title": "InjecAgent: Benchmarking Indirect Prompt Injections in Tool-Integrated LLM Agents (Zhan et al., 2024)",
        "link": "https://github.com/uiuc-kang-lab/InjecAgent",
        "description": "First indirect injection benchmark specifically for tool-integrated agents. Evaluates private-data exfiltration risk with explicit user-tool vs attacker-tool distinction across direct harm and data stealing categories.",
        "metrics": "Attack Success Rate (ASR) across tools, agents, harm categories, and exfiltration scenarios.",
        "baselines": "Multiple agent configurations evaluated. Used as evaluation benchmark by MindGuard (attack payload adaptation), Adaptive Attacks (100-case subset), and referenced by TraceAegis.",
        "how_to_use": "Clone the InjecAgent repository. Configure 17 user tools and 62 attacker tools. Run 1,054 test cases against target agents. Measure ASR per harm category. Requires Python and LLM API access.",
        "project_value": "The explicit exfiltration category is directly relevant to MCP data leakage risk assessment. The user-tool vs attacker-tool framework maps to the risk scorer's need to evaluate tool trustworthiness.",
    },
    {
        "benchmark_name": "R-Judge",
        "paper_title": "R-Judge: Benchmarking Safety Risk Awareness for LLM Agents (Yuan et al., 2024)",
        "link": "https://github.com/Lordog/R-Judge",
        "description": "Benchmark for evaluating risk awareness in LLM agent behaviors. Contains 569 curated multi-turn interaction records across 27 risk scenarios and 10 risk types sourced from WebArena, ToolEmu, InterCode-Bash/SQL, and MINT environments.",
        "metrics": "Safety judgment accuracy (binary safe/unsafe classification). Best: GPT-4o at 74.42%, most models near random.",
        "baselines": "11 LLMs evaluated with both prompting and fine-tuning approaches. Fine-tuning on safety judgment significantly improves performance.",
        "how_to_use": "Clone the R-Judge repository. Load the 569 interaction records. Evaluate target LLMs on safety judgment task. Compare against published baselines. Fine-tuning variant requires labeled training data and GPU resources.",
        "project_value": "The 10 risk types and 27 scenarios provide a template for MCP-specific risk categories. The finding that most models perform near random on risk judgment highlights the challenge the MCP risk scorer must overcome.",
    },
    {
        "benchmark_name": "DecodingTrust",
        "paper_title": "DecodingTrust: A Comprehensive Assessment of Trustworthiness in GPT Models (Wang et al., 2023)",
        "link": "https://huggingface.co/datasets/AI-Secure/DecodingTrust",
        "description": "Eight-dimensional trustworthiness benchmark (NeurIPS 2023 Outstanding Paper) evaluating toxicity, stereotype bias, adversarial robustness, OOD robustness, adversarial demonstrations, privacy, machine ethics, and fairness.",
        "metrics": "Dimension-specific metrics across 8 trustworthiness dimensions. Uses sub-benchmarks: RealToxicityPrompts, BBQ, AdvGLUE, MMLU, ETHICS, and more.",
        "baselines": "GPT-4, GPT-3.5. Key finding: GPT-4 more vulnerable to jailbreaking than GPT-3.5 despite better baselines.",
        "how_to_use": "Access dataset from HuggingFace (AI-Secure/DecodingTrust). Run evaluations using the 8-dimension framework. Each dimension has specific evaluation protocols and metrics. Requires LLM API access and the DecodingTrust evaluation codebase.",
        "project_value": "The 8-dimension framework provides a scoring rubric adaptable for MCP agent trustworthiness evaluation. Finding that capability ≠ trustworthiness is critical for MCP risk assessment design.",
    },
    {
        "benchmark_name": "TrustLLM",
        "paper_title": "TrustLLM: Trustworthiness in Large Language Models (Huang et al., 2024)",
        "link": "https://github.com/HowieHwong/TrustLLM",
        "description": "Largest-scale LLM trustworthiness benchmark (ICML 2024) evaluating 16 LLMs across 6 dimensions (truthfulness, safety, fairness, robustness, privacy, machine ethics) using 30+ datasets and 18 subcategories.",
        "metrics": "Dimension-specific scores across 6 dimensions, 18 subcategories. Uses 30+ public datasets for evaluation.",
        "baselines": "16 mainstream LLMs. Finding: proprietary LLMs generally more trustworthy; some over-optimize safety at cost of utility.",
        "how_to_use": "Clone the TrustLLM GitHub repository. Set up evaluation environment. Run target LLMs against 30+ datasets across 6 dimensions. Compare against 16-model baseline. Requires significant compute and API access for full evaluation.",
        "project_value": "The 6 trustworthiness dimensions can be adapted as risk scoring dimensions for MCP agents. The safety-utility tradeoff finding directly informs risk threshold calibration.",
    },
    {
        "benchmark_name": "CVSS v3.1 / NVD",
        "paper_title": "From Description to Score: Can LLMs Quantify Vulnerabilities? (Jafarikhah et al., 2026)",
        "link": "https://nvd.nist.gov/",
        "description": "The Common Vulnerability Scoring System (CVSS v3.1) with the National Vulnerability Database (NVD), used to evaluate whether LLMs can automatically score vulnerabilities from textual descriptions. Tests automated severity quantification across 31,000+ CVE entries.",
        "metrics": "Accuracy, Precision, Recall, F1, MAE across 8 CVSS base metrics. Severity scale: None/Low/Medium/High/Critical.",
        "baselines": "6 LLMs: GPT-4o, GPT-5, Llama-3.3-70B, Gemini-2.5-Flash, DeepSeek-R1, Grok-3. GPT-5 highest precision. Two-shot prompting optimal.",
        "how_to_use": "Access NVD data via API (https://nvd.nist.gov/). Extract CVE descriptions and CVSS scores. Implement LLM-based scoring with zero-shot to ten-shot prompting. Train meta-classifier combining LLM outputs. Requires LLM API access.",
        "project_value": "CVSS scoring methodology is directly transferable to MCP risk scoring. The text-to-severity-score approach mirrors how MCP Security could score tool descriptions. The meta-classifier ensemble approach applies to combining multiple risk signals.",
    },
    {
        "benchmark_name": "ASB (Agent Security Bench)",
        "paper_title": "Agent Security Bench (Zhang et al., 2025)",
        "link": "https://github.com/agiresearch/ASB",
        "description": "Benchmark for formalizing attacks and defenses in LLM-based agents with 6 attack prompt types. Used by both Progent and ToolSafe for standardized evaluation.",
        "metrics": "Utility (%), Attack Success Rate (ASR %) per attack type.",
        "baselines": "6 attack types: combined_attack, context_ignoring, escape_characters, fake_completion, naive, average. Defense baselines: delimiters, sandwich, instructional prevention.",
        "how_to_use": "Clone the ASB repository. Configure agent environments. Run all 6 attack types against target defenses. Measure utility-ASR tradeoff. Requires Python and LLM API access.",
        "project_value": "The 6 attack types provide input-level risk classification categories for the MCP risk scorer. Standardized evaluation enables comparison with Progent and ToolSafe results.",
    },
    {
        "benchmark_name": "AgentHarm",
        "paper_title": "AgentHarm: Harmful Agent Behavior Benchmark (Gray Swan / UK AISI, 2025)",
        "link": "https://huggingface.co/datasets/ai-safety-institute/AgentHarm",
        "description": "Benchmark for evaluating harmful agent behaviors including deepfakes, misinformation generation, and other dangerous capabilities. Published by UK AI Safety Institute and Gray Swan.",
        "metrics": "Attack Success Rate (ASR), Safety Rate, Helpfulness.",
        "baselines": "Multiple LLM agents evaluated on harmful behavior compliance rates.",
        "how_to_use": "Access dataset from HuggingFace (ai-safety-institute/AgentHarm). Evaluate target agents on harmful behavior requests. Measure compliance rates. Requires LLM API access.",
        "project_value": "Defines the highest severity tier for risk scoring — requests that should always be denied. Essential for calibrating the 'critical risk' threshold of the MCP risk scorer.",
    },
    {
        "benchmark_name": "HarmBench",
        "paper_title": "Referenced by TRiSM for Agentic AI (Raza et al., 2025) and multiple survey papers",
        "link": "https://github.com/centerforaisafety/HarmBench",
        "description": "Standardized benchmark for prompt injection and jailbreak evaluation from the Center for AI Safety. Widely referenced across MCP security literature for measuring attack robustness.",
        "metrics": "Attack Success Rate (ASR), Robustness scores.",
        "baselines": "Multiple jailbreak techniques and defense mechanisms evaluated.",
        "how_to_use": "Clone from GitHub. Run jailbreak attacks against target models. Measure ASR and robustness. Requires Python and LLM API access.",
        "project_value": "Standard jailbreak evaluation relevant to MCP scenarios where agents attempt to bypass tool restrictions through prompt manipulation.",
    },
    {
        "benchmark_name": "JailbreakBench",
        "paper_title": "Referenced by TRiSM for Agentic AI (Raza et al., 2025)",
        "link": "https://github.com/JailbreakBench/jailbreakbench",
        "description": "Benchmark focused specifically on jailbreak scenarios for LLMs. Referenced alongside HarmBench for comprehensive jailbreak evaluation.",
        "metrics": "Attack Success Rate (ASR), Robustness under jailbreak scenarios.",
        "baselines": "Multiple jailbreak techniques cataloged and evaluated.",
        "how_to_use": "Clone from GitHub. Configure jailbreak evaluation. Run against target LLMs. Requires Python and LLM API access.",
        "project_value": "Complements HarmBench for jailbreak-specific evaluation. Relevant to testing whether the MCP risk scorer can detect jailbreak attempts targeting tool access permissions.",
    },
    {
        "benchmark_name": "BIPIA (Benchmark for Indirect Prompt Injection Attacks)",
        "paper_title": "Referenced by Beyond the Protocol (Song et al., 2025) and multiple papers",
        "link": "https://github.com/microsoft/BIPIA",
        "description": "The first indirect prompt injection attack benchmark, developed by Microsoft. Referenced across MCP security literature as a foundational PI evaluation framework.",
        "metrics": "Attack success rates for indirect prompt injection. Detection accuracy.",
        "baselines": "Multiple LLM models and defense mechanisms evaluated.",
        "how_to_use": "Clone from GitHub (Microsoft). Run indirect PI attacks against target systems. Measure detection rates. Requires Python and LLM access.",
        "project_value": "Foundational PI benchmark applicable to MCP scenarios where external content injections reach the agent through tool responses. Tests the risk scorer's ability to detect indirect attacks.",
    },
    {
        "benchmark_name": "HELM (Holistic Evaluation of Language Models)",
        "paper_title": "Referenced by TRiSM for Agentic AI (Raza et al., 2025)",
        "link": "https://github.com/stanford-crfm/helm",
        "description": "Comprehensive LLM evaluation framework from Stanford CRFM covering robustness, fairness, and bias. Referenced in MCP security literature as a comprehensive model evaluation standard.",
        "metrics": "Comprehensive metrics across multiple dimensions: robustness, fairness, bias, accuracy.",
        "baselines": "Multiple mainstream LLMs evaluated across all dimensions.",
        "how_to_use": "Install HELM framework. Configure evaluation scenarios. Run comprehensive model evaluation. Requires significant compute resources.",
        "project_value": "Provides the broadest LLM evaluation methodology that can inform how to holistically assess MCP agent capabilities and risks across multiple dimensions.",
    },
    {
        "benchmark_name": "BFCL-v3 (Berkeley Function Calling Leaderboard)",
        "paper_title": "Used by MCIP (Jing et al., 2025) for utility evaluation",
        "link": "https://github.com/ShishirPatil/gorilla/tree/main/berkeley-function-call-leaderboard",
        "description": "The standard benchmark for evaluating LLM function calling capabilities. Used to measure whether security-focused MCP designs degrade practical function calling utility.",
        "metrics": "Overall function calling accuracy (%). Measures correct tool selection, parameter extraction, and execution.",
        "baselines": "Broad LLM comparison available on the public leaderboard (gorilla.cs.berkeley.edu).",
        "how_to_use": "Access via GitHub or HuggingFace. Run target LLMs on function calling tasks. Compare against leaderboard. Requires LLM API access.",
        "project_value": "Essential for measuring the utility cost of MCP risk scoring. A risk scorer that significantly degrades function calling accuracy is impractical. BFCL-v3 provides the accepted measurement standard.",
    },
    {
        "benchmark_name": "WebArena",
        "paper_title": "Referenced by R-Judge (Yuan et al., 2024) and TRiSM (Raza et al., 2025)",
        "link": "https://github.com/web-arena-x/webarena",
        "description": "Web navigation benchmark used as a source environment for R-Judge safety evaluation scenarios. Tests agent interactions in realistic web environments with potential security implications.",
        "metrics": "Task success rate, safety compliance in web navigation tasks.",
        "baselines": "Multiple LLM-based web agents evaluated.",
        "how_to_use": "Set up WebArena environment with web applications. Deploy agent in navigation tasks. Evaluate task success and safety. Requires Docker and compute resources.",
        "project_value": "Web navigation is a common MCP tool category. WebArena scenarios test whether the risk scorer correctly assesses web-interaction risks including content injection and data exposure.",
    },
    {
        "benchmark_name": "AgentBench",
        "paper_title": "Referenced by TRiSM (Raza et al., 2025) and SentinelAgent (He et al., 2025)",
        "link": "https://github.com/THUDM/AgentBench",
        "description": "Multi-task agent benchmark for evaluating LLM agents across diverse environments. Referenced in MCP security literature for comprehensive agent capability assessment.",
        "metrics": "Task success rate, compliance metrics, multi-step attack resilience.",
        "baselines": "Multiple LLM agents across diverse task categories.",
        "how_to_use": "Clone from GitHub. Set up multi-task evaluation environments. Run agent evaluations. Requires Python and LLM API access.",
        "project_value": "Provides multi-task evaluation methodology applicable to testing the MCP risk scorer across diverse tool-use scenarios. Multi-step evaluation tests sustained risk assessment accuracy.",
    },
]


# ---------------------------------------------------------------------------
# SECTION 3: MARKDOWN GENERATORS
# ---------------------------------------------------------------------------

def generate_datasets_markdown() -> str:
    lines = [
        "# Datasets Review — MCP Security Literature",
        "",
        f"> Comprehensive review of **{len(DATASETS)} datasets** extracted from 82 papers",
        "> across the MCP Security literature review.  ",
        f"> Generated: {TODAY}",
        "",
        "---",
        "",
        "## Table of Contents",
        "",
    ]
    for i, ds in enumerate(DATASETS, 1):
        anchor = ds["dataset_name"].lower().replace(" ", "-").replace("(", "").replace(")", "").replace("/", "").replace(",", "")
        lines.append(f"{i}. [{ds['dataset_name']}](#{i}-{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    for i, ds in enumerate(DATASETS, 1):
        lines.append(f"## {i}. {ds['dataset_name']}")
        lines.append("")
        lines.append(f"**Paper:** {ds['paper_title']}")
        lines.append("")
        lines.append("| Field | Details |")
        lines.append("|-------|---------|")
        lines.append(f"| **Direct Link** | {ds['link']} |")
        lines.append(f"| **Description** | {ds['description']} |")
        lines.append(f"| **Structure Details** | {ds['structure']} |")
        lines.append(f"| **How the Paper Used It** | {ds['paper_usage']} |")
        lines.append(f"| **How It Can Help My Project** | {ds['project_value']} |")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def generate_benchmarks_markdown() -> str:
    lines = [
        "# Benchmarks Review — MCP Security Literature",
        "",
        f"> Comprehensive review of **{len(BENCHMARKS)} benchmarks** extracted from 82 papers",
        "> across the MCP Security literature review.  ",
        f"> Generated: {TODAY}",
        "",
        "---",
        "",
        "## Table of Contents",
        "",
    ]
    for i, bm in enumerate(BENCHMARKS, 1):
        anchor = bm["benchmark_name"].lower().replace(" ", "-").replace("(", "").replace(")", "").replace("/", "").replace(",", "")
        lines.append(f"{i}. [{bm['benchmark_name']}](#{i}-{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    for i, bm in enumerate(BENCHMARKS, 1):
        lines.append(f"## {i}. {bm['benchmark_name']}")
        lines.append("")
        lines.append(f"**Paper:** {bm['paper_title']}")
        lines.append("")
        lines.append("| Field | Details |")
        lines.append("|-------|---------|")
        lines.append(f"| **Link** | {bm['link']} |")
        lines.append(f"| **Description** | {bm['description']} |")
        lines.append(f"| **Metrics Used** | {bm['metrics']} |")
        lines.append(f"| **Common Baselines** | {bm['baselines']} |")
        lines.append(f"| **How to Use It** | {bm['how_to_use']} |")
        lines.append(f"| **How It Can Help My Project** | {bm['project_value']} |")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# SECTION 4: WORD (.docx) GENERATORS
# ---------------------------------------------------------------------------

def _add_styled_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (field, value) in enumerate(rows):
        cell_field = table.rows[idx].cells[0]
        cell_value = table.rows[idx].cells[1]
        run_f = cell_field.paragraphs[0].add_run(field)
        run_f.bold = True
        run_f.font.size = Pt(10)
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(10)
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)


def generate_datasets_docx(path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Datasets Review — MCP Security Literature", level=0)
    doc.add_paragraph(
        f"Comprehensive review of {len(DATASETS)} datasets extracted from 82 papers "
        f"across the MCP Security literature review. Generated: {TODAY}"
    )

    for i, ds in enumerate(DATASETS, 1):
        doc.add_heading(f"{i}. {ds['dataset_name']}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"Paper: {ds['paper_title']}")
        run.italic = True
        run.font.size = Pt(10)

        _add_styled_table(doc, [
            ("Direct Link", ds["link"]),
            ("Description", ds["description"]),
            ("Structure Details", ds["structure"]),
            ("How the Paper Used It", ds["paper_usage"]),
            ("How It Can Help My Project", ds["project_value"]),
        ])
        doc.add_paragraph("")  # spacer

    doc.save(str(path))


def generate_benchmarks_docx(path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Benchmarks Review — MCP Security Literature", level=0)
    doc.add_paragraph(
        f"Comprehensive review of {len(BENCHMARKS)} benchmarks extracted from 82 papers "
        f"across the MCP Security literature review. Generated: {TODAY}"
    )

    for i, bm in enumerate(BENCHMARKS, 1):
        doc.add_heading(f"{i}. {bm['benchmark_name']}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"Paper: {bm['paper_title']}")
        run.italic = True
        run.font.size = Pt(10)

        _add_styled_table(doc, [
            ("Link", bm["link"]),
            ("Description", bm["description"]),
            ("Metrics Used", bm["metrics"]),
            ("Common Baselines", bm["baselines"]),
            ("How to Use It", bm["how_to_use"]),
            ("How It Can Help My Project", bm["project_value"]),
        ])
        doc.add_paragraph("")  # spacer

    doc.save(str(path))


# ---------------------------------------------------------------------------
# SECTION 5: EXCEL (.xlsx) GENERATORS
# ---------------------------------------------------------------------------

HEADER_FONT = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
CELL_ALIGN = Alignment(vertical="top", wrap_text=True)
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)
ALT_FILL = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")


def _style_header(ws, headers: list[str], widths: list[int]) -> None:
    for col_idx, (header, width) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = HEADER_ALIGN
        cell.border = THIN_BORDER
        ws.column_dimensions[cell.column_letter].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def generate_datasets_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Datasets"

    headers = [
        "Dataset Name", "Paper Title", "Direct Link",
        "Description", "Structure Details",
        "How the Paper Used It", "How It Can Help My Project",
    ]
    widths = [30, 50, 45, 55, 50, 50, 55]
    _style_header(ws, headers, widths)

    for row_idx, ds in enumerate(DATASETS, 2):
        values = [
            ds["dataset_name"], ds["paper_title"], ds["link"],
            ds["description"], ds["structure"],
            ds["paper_usage"], ds["project_value"],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = CELL_ALIGN
            cell.border = THIN_BORDER
            cell.font = Font(size=10, name="Calibri")
            if row_idx % 2 == 0:
                cell.fill = ALT_FILL

    wb.save(str(path))


def generate_benchmarks_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Benchmarks"

    headers = [
        "Benchmark Name", "Paper Title", "Link",
        "Description", "Metrics Used", "Common Baselines",
        "How to Use It", "How It Can Help My Project",
    ]
    widths = [28, 50, 45, 55, 45, 50, 55, 55]
    _style_header(ws, headers, widths)

    for row_idx, bm in enumerate(BENCHMARKS, 2):
        values = [
            bm["benchmark_name"], bm["paper_title"], bm["link"],
            bm["description"], bm["metrics"], bm["baselines"],
            bm["how_to_use"], bm["project_value"],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = CELL_ALIGN
            cell.border = THIN_BORDER
            cell.font = Font(size=10, name="Calibri")
            if row_idx % 2 == 0:
                cell.fill = ALT_FILL

    wb.save(str(path))


# ---------------------------------------------------------------------------
# SECTION 6: MAIN
# ---------------------------------------------------------------------------

def main() -> None:
    print(f"Generating review files in {OUTPUT_DIR} ...")

    # Markdown
    md_ds = OUTPUT_DIR / "datasets_review.md"
    md_ds.write_text(generate_datasets_markdown(), encoding="utf-8")
    print(f"  Created {md_ds.name} ({len(DATASETS)} datasets)")

    md_bm = OUTPUT_DIR / "benchmarks_review.md"
    md_bm.write_text(generate_benchmarks_markdown(), encoding="utf-8")
    print(f"  Created {md_bm.name} ({len(BENCHMARKS)} benchmarks)")

    # Word
    docx_ds = OUTPUT_DIR / "datasets_review.docx"
    generate_datasets_docx(docx_ds)
    print(f"  Created {docx_ds.name}")

    docx_bm = OUTPUT_DIR / "benchmarks_review.docx"
    generate_benchmarks_docx(docx_bm)
    print(f"  Created {docx_bm.name}")

    # Excel
    xlsx_ds = OUTPUT_DIR / "datasets_review.xlsx"
    generate_datasets_xlsx(xlsx_ds)
    print(f"  Created {xlsx_ds.name}")

    xlsx_bm = OUTPUT_DIR / "benchmarks_review.xlsx"
    generate_benchmarks_xlsx(xlsx_bm)
    print(f"  Created {xlsx_bm.name}")

    print(f"\nDone. {len(DATASETS)} datasets + {len(BENCHMARKS)} benchmarks across 6 files.")


if __name__ == "__main__":
    main()
